from flask import Flask, request, jsonify
from flask_cors import CORS
import os
import pandas as pd
from docx import Document

app = Flask(__name__)
CORS(app)  # Adicione esta linha para habilitar CORS

def salvar_ministerios_em_arquivos(df, diretorio_saida):
    if not os.path.exists(diretorio_saida):
        os.makedirs(diretorio_saida)
    colunas_ministerio = ['Ministério 1', 'Ministério 2', 'Ministério 3']
    for coluna in colunas_ministerio:
        df[coluna] = df[coluna].str.strip()
        valores_unicos_ministerio = df[coluna].dropna().unique()
        for ministerio in valores_unicos_ministerio:
            ministerio_df = df[df[coluna] == ministerio]
            caminho_arquivo = os.path.join(diretorio_saida, f'{ministerio}.xlsx')
            ministerio_df.to_excel(caminho_arquivo, index=False)

def criar_documento_direcionamento(df, ministerio, caminho, diretorio_saida, mes):
    if pd.notnull(ministerio):
        diretorio_saida_docx = os.path.join(diretorio_saida, 'docx')
        if not os.path.exists(diretorio_saida_docx):
            os.makedirs(diretorio_saida_docx)
        documento = Document()
        for linha in df.index:
            nome = df.loc[linha, 'Nome']
            telefone = df.loc[linha, 'Contato']
            apto = df.loc[linha, 'APTO P/ SERVIR']
            faltaum = 'X' if pd.isnull(df.loc[linha, 1]) else ''
            faltadois = 'X' if pd.isnull(df.loc[linha, 2]) else ''
            faltatres = 'X' if pd.isnull(df.loc[linha, 3]) else ''
            faltaquatro = 'X' if pd.isnull(df.loc[linha, 4]) else ''
            ficha = df.loc[linha, 'FICHA']
            quatrosim = 'X' if df.loc[linha, 4] == 'Sim' else ''
            quatronao = 'X' if df.loc[linha, 4] == 'Não' else ''
            nove = 'X' if str(df.loc[linha, 'H']).strip() in ['9', '9h'] else ''
            seis = 'X' if str(df.loc[linha, 'H']).strip() in ['18', '18h'] else ''
            minis = df.loc[linha, 'Ministério 1']
            mes2 = df.loc[linha, 'Unnamed: 4']

            referencias = {
                "AAAA": nome,
                "BBBB": str(telefone),
                "CC": apto,
                "DD": faltaum,
                "EE": faltadois,
                "FF": faltatres,
                "GG": faltaquatro,
                "HH": ficha,
                "II": quatrosim,
                "JJ": quatronao,
                "KK": nove,
                "LL": seis,
                "XX": minis
            }

            for paragrafo in Document(caminho).paragraphs:
                novo_paragrafo = paragrafo.text
                for codigo, valor in referencias.items():
                    novo_paragrafo = novo_paragrafo.replace(codigo, str(valor))
                if mes2 == mes:
                    documento.add_paragraph(novo_paragrafo)

        caminho_completo = os.path.join(diretorio_saida_docx, f"{ministerio}_documento.docx")
        documento.save(caminho_completo)

@app.route('/generateDocs', methods=['POST'])
def gerar_documentos():
    data = request.json
    mes = data.get('mes')
    print(f"Received month: {mes}")

    caminho_do_arquivo = r'C:\Users\gabri\OneDrive\Área de Trabalho\ondas-py-back\CHAMADACADASTRO3.xlsx'
    diretorio_saida = r'C:\Users\gabri\OneDrive\Área de Trabalho\ondas-py-back\ministerios'
    df = pd.read_excel(caminho_do_arquivo)
    df = df.iloc[2:]
    salvar_ministerios_em_arquivos(df, diretorio_saida)
    caminho_arquivo_word = r'C:\Users\gabri\OneDrive\Área de Trabalho\ondas-py-back\Direcionamentos.docx'

    for ministerio in df['Ministério 1'].unique():
        if pd.notnull(ministerio):
            criar_documento_direcionamento(df[df['Ministério 1'] == ministerio], ministerio, caminho_arquivo_word, diretorio_saida, mes)

    return jsonify({'message': 'Documentos gerados com sucesso!'})

@app.route('/')
def index():
    return jsonify({'message': 'API is running'})

if __name__ == '__main__':
    app.run(debug=True)